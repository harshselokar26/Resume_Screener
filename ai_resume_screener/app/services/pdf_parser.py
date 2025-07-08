"""
PDF Parser Service for AI Resume Screener

This module handles extraction of text content from PDF and DOC/DOCX files
using various parsing libraries with fallback mechanisms.
"""

import os
import logging
from typing import Optional, Dict, Any
from pathlib import Path
import asyncio
from concurrent.futures import ThreadPoolExecutor

# PDF parsing libraries
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from pdfminer.high_level import extract_text as pdfminer_extract
    from pdfminer.layout import LAParams
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

# DOC/DOCX parsing libraries
try:
    import docx
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import win32com.client
    WIN32COM_AVAILABLE = True
except ImportError:
    WIN32COM_AVAILABLE = False

from app.config.settings import settings
from app.utils.exceptions import FileProcessingError

# Setup logging
logger = logging.getLogger(__name__)


class PDFParser:
    """
    PDF and document parser with multiple parsing strategies.
    """
    
    def __init__(self):
        """Initialize PDF parser with available libraries."""
        self.executor = ThreadPoolExecutor(max_workers=2)
        self._log_available_parsers()
    
    def _log_available_parsers(self):
        """Log available parsing libraries."""
        available_parsers = []
        if PYMUPDF_AVAILABLE:
            available_parsers.append("PyMuPDF")
        if PDFMINER_AVAILABLE:
            available_parsers.append("pdfminer")
        if PYTHON_DOCX_AVAILABLE:
            available_parsers.append("python-docx")
        if WIN32COM_AVAILABLE:
            available_parsers.append("win32com")
        
        logger.info(f"Available parsers: {', '.join(available_parsers)}")
    
    async def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file using available parsers.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            str: Extracted text content
            
        Raises:
            FileProcessingError: If text extraction fails
        """
        if not os.path.exists(file_path):
            raise FileProcessingError(f"File not found: {file_path}")
        
        try:
            # Try PyMuPDF first (fastest and most reliable)
            if PYMUPDF_AVAILABLE:
                text = await self._extract_with_pymupdf(file_path)
                if text and text.strip():
                    logger.info(f"Successfully extracted text using PyMuPDF: {len(text)} chars")
                    return text
            
            # Fallback to pdfminer
            if PDFMINER_AVAILABLE:
                text = await self._extract_with_pdfminer(file_path)
                if text and text.strip():
                    logger.info(f"Successfully extracted text using pdfminer: {len(text)} chars")
                    return text
            
            raise FileProcessingError("No PDF parsing libraries available")
            
        except Exception as e:
            logger.error(f"PDF text extraction failed: {str(e)}")
            raise FileProcessingError(f"Failed to extract text from PDF: {str(e)}")
    
    async def _extract_with_pymupdf(self, file_path: str) -> str:
        """Extract text using PyMuPDF."""
        def _extract():
            doc = fitz.open(file_path)
            text = ""
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text += page.get_text()
            doc.close()
            return text
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, _extract)
    
    async def _extract_with_pdfminer(self, file_path: str) -> str:
        """Extract text using pdfminer."""
        def _extract():
            laparams = LAParams(
                boxes_flow=0.5,
                word_margin=0.1,
                char_margin=2.0,
                line_margin=0.5
            )
            return pdfminer_extract(file_path, laparams=laparams)
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, _extract)
    
    async def extract_text_from_doc(self, file_path: str) -> str:
        """
        Extract text from DOC/DOCX files.
        
        Args:
            file_path: Path to DOC/DOCX file
            
        Returns:
            str: Extracted text content
            
        Raises:
            FileProcessingError: If text extraction fails
        """
        if not os.path.exists(file_path):
            raise FileProcessingError(f"File not found: {file_path}")
        
        file_extension = Path(file_path).suffix.lower()
        
        try:
            if file_extension == '.docx':
                return await self._extract_from_docx(file_path)
            elif file_extension == '.doc':
                return await self._extract_from_doc(file_path)
            else:
                raise FileProcessingError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"DOC/DOCX text extraction failed: {str(e)}")
            raise FileProcessingError(f"Failed to extract text from document: {str(e)}")
    
    async def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        if not PYTHON_DOCX_AVAILABLE:
            raise FileProcessingError("python-docx library not available")
        
        def _extract():
            doc = docx.Document(file_path)
            text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
            
            return '\n'.join(text)
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, _extract)
    
    async def _extract_from_doc(self, file_path: str) -> str:
        """Extract text from DOC file."""
        if WIN32COM_AVAILABLE:
            return await self._extract_doc_with_win32com(file_path)
        else:
            # Try converting to DOCX first
            raise FileProcessingError(
                "DOC file support requires win32com library or manual conversion to DOCX"
            )
    
    async def _extract_doc_with_win32com(self, file_path: str) -> str:
        """Extract text from DOC file using win32com (Windows only)."""
        def _extract():
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(file_path)
            text = doc.Content.Text
            doc.Close()
            word.Quit()
            return text
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, _extract)
    
    async def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dict containing file metadata
        """
        metadata = {
            "file_size": 0,
            "page_count": 0,
            "creation_date": None,
            "modification_date": None,
            "author": None,
            "title": None,
            "subject": None
        }
        
        try:
            # Get file stats
            file_stats = os.stat(file_path)
            metadata["file_size"] = file_stats.st_size
            
            # Extract PDF metadata if available
            if PYMUPDF_AVAILABLE and file_path.lower().endswith('.pdf'):
                def _extract_metadata():
                    doc = fitz.open(file_path)
                    pdf_metadata = doc.metadata
                    metadata.update({
                        "page_count": doc.page_count,
                        "author": pdf_metadata.get("author"),
                        "title": pdf_metadata.get("title"),
                        "subject": pdf_metadata.get("subject"),
                        "creation_date": pdf_metadata.get("creationDate"),
                        "modification_date": pdf_metadata.get("modDate")
                    })
                    doc.close()
                    return metadata
                
                loop = asyncio.get_event_loop()
                metadata = await loop.run_in_executor(self.executor, _extract_metadata)
            
        except Exception as e:
            logger.warning(f"Failed to extract metadata: {str(e)}")
        
        return metadata
    
    async def validate_file(self, file_path: str) -> bool:
        """
        Validate if file can be processed.
        
        Args:
            file_path: Path to file
            
        Returns:
            bool: True if file is valid and can be processed
        """
        try:
            if not os.path.exists(file_path):
                return False
            
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.pdf':
                return await self._validate_pdf(file_path)
            elif file_extension in ['.doc', '.docx']:
                return await self._validate_doc(file_path)
            
            return False
            
        except Exception as e:
            logger.error(f"File validation failed: {str(e)}")
            return False
    
    async def _validate_pdf(self, file_path: str) -> bool:
        """Validate PDF file."""
        try:
            if PYMUPDF_AVAILABLE:
                def _validate():
                    doc = fitz.open(file_path)
                    is_valid = doc.page_count > 0
                    doc.close()
                    return is_valid
                
                loop = asyncio.get_event_loop()
                return await loop.run_in_executor(self.executor, _validate)
            
            return True  # Assume valid if we can't check
            
        except Exception:
            return False
    
    async def _validate_doc(self, file_path: str) -> bool:
        """Validate DOC/DOCX file."""
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.docx' and PYTHON_DOCX_AVAILABLE:
                def _validate():
                    doc = docx.Document(file_path)
                    return len(doc.paragraphs) > 0 or len(doc.tables) > 0
                
                loop = asyncio.get_event_loop()
                return await loop.run_in_executor(self.executor, _validate)
            
            return True  # Assume valid if we can't check
            
        except Exception:
            return False
    
    async def health_check(self) -> bool:
        """
        Check if PDF parser service is healthy.
        
        Returns:
            bool: True if service is healthy
        """
        try:
            # Check if at least one parser is available
            return PYMUPDF_AVAILABLE or PDFMINER_AVAILABLE
        except Exception:
            return False
    
    def get_supported_formats(self) -> list:
        """
        Get list of supported file formats.
        
        Returns:
            list: Supported file extensions
        """
        formats = []
        
        if PYMUPDF_AVAILABLE or PDFMINER_AVAILABLE:
            formats.append('.pdf')
        
        if PYTHON_DOCX_AVAILABLE:
            formats.append('.docx')
        
        if WIN32COM_AVAILABLE:
            formats.append('.doc')
        
        return formats
